from django.shortcuts import render
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Proposal, ProposalHistory
from .serializers import ProposalSerializer, ProposalHistorySerializer
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from reportlab.pdfgen import canvas
from io import BytesIO
import os
from django.conf import settings
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from datetime import datetime

def HEXtoRGB(hex_color):
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

class ProposalViewSet(viewsets.ModelViewSet):
    queryset = Proposal.objects.all()
    serializer_class = ProposalSerializer

    def perform_create(self, serializer):
        user_ip = self.request.META.get('REMOTE_ADDR')
        serializer.save(user_ip=user_ip)

    def perform_update(self, serializer):
        user_ip = self.request.META.get('REMOTE_ADDR')
        serializer.save(user_ip=user_ip)

    @action(detail=True, methods=['post'], url_path='generate-document')
    def generate_document(self, request, pk=None):
        proposal = self.get_object()
        doc_type = request.data.get('doc_type', 'pdf')

        current_time = datetime.now().strftime("%d%m%Y_%H%M")
        date_generation = datetime.now().strftime("%d/%m/%Y")
        filename_base = f"Proposition_commerciale_{proposal.opportunity_number}_{current_time}"

        cost = proposal.ouvrage_cost if proposal.ouvrage_cost is not None else 0
        logo_path = os.path.join(settings.BASE_DIR, 'static', 'logo.png') 
        print(f"Corrected Logo path: {logo_path}")

        if doc_type == 'pdf':
            try:
                pdf_buffer = BytesIO()
                doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                                            topMargin=0.5*inch, bottomMargin=0.25*inch)
                styles = getSampleStyleSheet()
                story = []
                style_body = ParagraphStyle('BodyText', parent=styles['Normal'], fontSize=11, leading=12, spaceBefore=5, spaceAfter=5)
                style_body_blue = ParagraphStyle('BodyBlueText', parent=style_body, textColor="#00008f")
                style_main_title = ParagraphStyle('MainTitle', parent=styles['h1'], fontSize=11, leading=18, spaceBefore=6, spaceAfter=5, textColor="#00008f")
                style_section_title = ParagraphStyle('SectionTitle', parent=styles['h2'], fontSize=11, leading=13, fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=10)
                style_table_text = ParagraphStyle('TableText', parent=styles['Normal'], fontSize=9, leading=11)
                style_table_cell_text = ParagraphStyle('TableCellText', parent=style_table_text, alignment=TA_RIGHT)
                style_footnote = ParagraphStyle('Footnote', parent=styles['Normal'], fontSize=10, leading=10, spaceBefore=6)
                style_footnote_indented = ParagraphStyle('FootnoteIndented', parent=style_footnote, leftIndent=0.27*inch)
                style_body_html = ParagraphStyle('BodyTextHTML', parent=style_table_text, allowHTMLEscapes=1)

                if logo_path and os.path.exists(logo_path):
                    logo_img = Image(logo_path, width=0.75*inch, height=0.75*inch)
                    logo_img.hAlign = 'CENTER'
                    story.append(logo_img)
                else:
                    print(f"Logo PDF non trouvé ou chemin non défini.")
                
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph("TARIFICATION INDICATIVE", style_main_title))
                story.append(Paragraph("Produit chantier", style_body_blue))
                story.append(Paragraph("Tarification indicative (*) sur la base d'un risque conforme aux paramètres suivants :", style_body))

                params_data = [
                    (f"Type d'ouvrage:", proposal.get_ouvrage_destination_display() if proposal.ouvrage_destination else "Non Applicable"),
                    (f"Types de travaux réalisés:", proposal.get_work_type_display() if proposal.work_type else "rénovation légère, rénovation lourde, construction neuve"),
                    (f"Coût du chantier:", f"{cost:.2f} €" if proposal.ouvrage_cost is not None else "xxxxxxx €"),
                    (f"Présence d'existant:", "Oui" if proposal.existing_presence else "Non"),
                    (f"Garantie choisie:", proposal.get_guarantee_type_display() if proposal.guarantee_type else "Non renseigné"),
                    (f"Description de l'ouvrage :", proposal.ouvrage_description if proposal.ouvrage_description else "Non renseignée"),
                    (f"Adresse du chantier:", proposal.address_chantier if proposal.address_chantier else "Non renseignée")
                ]
                for key, value in params_data:
                    p_text = f"{key} {value}" 
                    story.append(Paragraph(p_text, style_body)) 

                story.append(Paragraph("<u>Garanties Tous Risques chantier</u>", style_body_html)) 
                story.append(Paragraph("MONTANTS DE GARANTIES (exprimés en €)", style_section_title))
                data_trc = [
                    [Paragraph("Dommages matériels à l'ouvrage<br/>(Ce montant est épuisable pendant la durée des travaux)", style_body_html), Paragraph("XXXXXXXX", style_table_cell_text)],
                    [Paragraph("Responsabilité civile (tous dommages confondus)<br/>(Ces montants sont épuisables pendant la durée des travaux)", style_body_html), Paragraph("XXXXXXXXX", style_table_cell_text)],
                    [Paragraph("Maintenance-visite<br/>(Ce montant est compris dans le montant de la garantie des dommages matériels à l'ouvrage)", style_body_html), Paragraph("XXXXXXXXXXX", style_table_cell_text)],
                    [Paragraph("Mesure conservatoire", style_table_text), Paragraph("XXXXXXXXXXXX", style_table_cell_text)]
                ]
                table_trc = Table(data_trc)
                table_trc.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.75, colors.black),
                    ('LEFTPADDING', (0,0), (-1,-1), 5),
                    ('RIGHTPADDING', (0,0), (-1,-1), 5),
                    ('TOPPADDING', (0,0), (-1,-1), 3),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 3),
                ]))
                story.append(table_trc)

                story.append(Paragraph("MONTANTS DE FRANCHISES (par sinistre exprimés en €)", style_section_title))
                data_franchises = [
                    [Paragraph("Dommages subis par les ouvrages de bâtiment", style_table_text), Paragraph("XXXXXXXXXXX", style_table_cell_text)],
                    [Paragraph("Catastrophes naturelles", style_table_text), Paragraph("Montant déterminé par la loi ou par ses textes d'application", style_table_cell_text)],
                    [Paragraph("Responsabilité civile (1)", style_body_html), ""],
                    [Paragraph("- Assuré maître d'ouvrage", style_table_text), Paragraph("XXXXXXXXXXX", style_table_cell_text)],
                    [Paragraph("- Assurés intervenants", style_table_text), Paragraph("SANS", style_table_cell_text)],
                    [Paragraph("Maintenance-visite", style_table_text), Paragraph("XXXXXXXXXXX", style_table_cell_text)]
                ]
                table_franchises = Table(data_franchises)
                table_franchises.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.75, colors.black),
                    ('LEFTPADDING', (0,0), (-1,-1), 5),
                    ('RIGHTPADDING', (0,0), (-1,-1), 5),
                    ('TOPPADDING', (0,0), (-1,-1), 3),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 3),
                    ('SPAN', (1,2), (1,2)),
                    ('LEFTPADDING', (0,3), (0,4), 20),
                ]))
                story.append(table_franchises)
                story.append(Paragraph("(1) Ces franchises s'appliquent pour des dommages autres que corporels", style_footnote_indented))
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(f"Date de simulation de tarif: le {date_generation}", style_body))
                story.append(Paragraph("(*) Cette tarification est faite sous réserve d'acceptation du risque par la compagnie", style_footnote))

                doc_pdf.build(story)
                pdf_value = pdf_buffer.getvalue()
                pdf_buffer.close()
                response = HttpResponse(pdf_value, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
                return response
            except Exception as e:
                return Response({'error': f"Erreur lors de la génération du PDF: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        elif doc_type == 'word':
            try:
                word_buffer = BytesIO()
                doc_word = DocxDocument()

                sections = doc_word.sections
                for section in sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.25)
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)

                # Fonctions d'aide pour formater le texte et les paragraphes
                def set_run_font(run, name='Calibri', size_pt=11, bold=False, italic=False, color_rgb=None, underline=False):
                    font = run.font
                    font.name = name
                    font.size = Pt(size_pt)
                    font.bold = bold
                    font.italic = italic
                    if color_rgb:
                        font.color.rgb = color_rgb
                    font.underline = underline

                def set_paragraph_format(paragraph, alignment=None, space_before_pt=None, space_after_pt=None, 
                                         line_spacing_rule=None, line_spacing_value=None, left_indent_inches=None, keep_with_next=False):
                    p_fmt = paragraph.paragraph_format
                    if alignment:
                        p_fmt.alignment = alignment
                    if space_before_pt:
                        p_fmt.space_before = Pt(space_before_pt)
                    if space_after_pt:
                        p_fmt.space_after = Pt(space_after_pt)
                    if line_spacing_rule: 
                        p_fmt.line_spacing_rule = line_spacing_rule
                    if line_spacing_value:
                        p_fmt.line_spacing = line_spacing_value
                    if left_indent_inches:
                        p_fmt.left_indent = Inches(left_indent_inches)
                    if keep_with_next:
                        p_fmt.keep_with_next = True
                
                style_normal_word = doc_word.styles['Normal']
                font_normal = style_normal_word.font
                font_normal.name = 'Calibri'
                font_normal.size = Pt(11)
                p_fmt_normal = style_normal_word.paragraph_format
                p_fmt_normal.space_before = Pt(5)
                p_fmt_normal.space_after = Pt(5)
                p_fmt_normal.line_spacing_rule = None
                p_fmt_normal.line_spacing = 1

                if logo_path and os.path.exists(logo_path):
                    p_logo = doc_word.add_paragraph()
                    set_paragraph_format(p_logo, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=0, space_after_pt=0)
                    run_logo = p_logo.add_run()
                    run_logo.add_picture(logo_path, width=Inches(0.75))
                else:
                    print(f"Logo Word non trouvé ou chemin non défini.")


                p_main_title = doc_word.add_paragraph()
                set_paragraph_format(p_main_title, alignment=WD_ALIGN_PARAGRAPH.LEFT, 
                                     space_before_pt=6, space_after_pt=5, 
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1)
                run_main_title = p_main_title.add_run("TARIFICATION INDICATIVE")
                set_run_font(run_main_title, name='Calibri', size_pt=13, bold=False, color_rgb=HEXtoRGB("#00008f"))


                p_subtitle = doc_word.add_paragraph()
                set_paragraph_format(p_subtitle, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     space_before_pt=5, space_after_pt=5, 
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1)
                run_subtitle = p_subtitle.add_run("Produit chantier")
                set_run_font(run_subtitle, name='Calibri', size_pt=11, color_rgb=HEXtoRGB("#00008f"))

                p_intro = doc_word.add_paragraph()
                set_paragraph_format(p_intro, space_before_pt=5, space_after_pt=5, 
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1)
                run_intro = p_intro.add_run("Tarification indicative (*) sur la base d'un risque conforme aux paramètres suivants :")
                set_run_font(run_intro, name='Calibri', size_pt=11)

                params_data_word = [
                    ("Type d'ouvrage:", proposal.get_ouvrage_destination_display() if proposal.ouvrage_destination else "Non Applicable"),
                    ("Types de travaux réalisés:", proposal.get_work_type_display() if proposal.work_type else "rénovation légère, rénovation lourde, construction neuve"),
                    ("Coût du chantier:", f"{cost:.2f} €" if proposal.ouvrage_cost is not None else "xxxxxxx €"),
                    ("Présence d'existant:", "Oui" if proposal.existing_presence else "Non"),
                    ("Garantie choisie:", proposal.get_guarantee_type_display() if proposal.guarantee_type else "DO/TRC/ DO + TRC"),
                    ("Description de l'ouvrage :", proposal.ouvrage_description if proposal.ouvrage_description else "Non renseignée"),
                    ("Adresse du chantier:", proposal.address_chantier if proposal.address_chantier else "Non renseignée")
                ]
                for key, value in params_data_word:
                    p_param = doc_word.add_paragraph()
                    set_paragraph_format(p_param, space_before_pt=5, space_after_pt=5, 
                                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1)
                    run_param_key = p_param.add_run(key + " ") 
                    set_run_font(run_param_key, name='Calibri', size_pt=11)
                    run_param_value = p_param.add_run(value)
                    set_run_font(run_param_value, name='Calibri', size_pt=11)

                p_garanties_text = doc_word.add_paragraph()
                set_paragraph_format(p_garanties_text, space_before_pt=5, space_after_pt=5, 
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1.22) 
                run_garanties_text = p_garanties_text.add_run("Garanties Tous Risques chantier")
                set_run_font(run_garanties_text, name='Calibri', size_pt=9, underline=True)
                
                p_montants_title = doc_word.add_paragraph()
                set_paragraph_format(p_montants_title, space_before_pt=10, space_after_pt=10,
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1.18) 
                run_montants_title = p_montants_title.add_run("MONTANTS DE GARANTIES (exprimés en €)")
                set_run_font(run_montants_title, name='Calibri', size_pt=11, bold=True)

                # Helper pour formater les cellules de tableau
                def format_table_cell(cell, text, font_name='Calibri', size_pt=9, bold=False, 
                                      alignment=None, left_indent_inches=None, is_html_like=False,
                                      cell_v_align=WD_CELL_VERTICAL_ALIGNMENT.TOP):
                    cell.vertical_alignment = cell_v_align
                    if cell.paragraphs and not cell.paragraphs[0].text.strip():
                         p = cell.paragraphs[0]
                         p.clear()
                         set_paragraph_format(p, space_before_pt=0, space_after_pt=0, line_spacing_rule=WD_LINE_SPACING.SINGLE)

                    else:
                        p = cell.add_paragraph()
                        set_paragraph_format(p, space_before_pt=0, space_after_pt=0, line_spacing_rule=WD_LINE_SPACING.SINGLE)


                    if alignment:
                        p.alignment = alignment
                    if left_indent_inches:
                        p.paragraph_format.left_indent = Inches(left_indent_inches)
                    
                    
                    lines = text.split('<br/>') if is_html_like and isinstance(text, str) else [text]
                    for i, line_text in enumerate(lines):
                        run = p.add_run(line_text)
                        set_run_font(run, name=font_name, size_pt=size_pt, bold=bold)
                        if i < len(lines) - 1:
                            run.add_break() 

                table_data_trc_word = [
                    ("Dommages matériels à l'ouvrage (Ce montant est épuisable pendant la durée des travaux)", "XXXXXXXX"),
                    ("Responsabilité civile (tous dommages confondus)<br/>(Ces montants sont épuisables pendant la durée des travaux)", "XXXXXXXXX"),
                    ("Maintenance-visite (Ce montant est compris dans le montant de la garantie des dommages matériels à l'ouvrage)", "XXXXXXXXXXX"),
                    ("Mesure conservatoire", "XXXXXXXXXXXX")
                ]
                table_trc_word = doc_word.add_table(rows=len(table_data_trc_word), cols=2)
                table_trc_word.style = 'TableGrid' 
                table_trc_word.autofit = True


                for i, (item, value) in enumerate(table_data_trc_word):
                    cell1 = table_trc_word.cell(i, 0)
                    format_table_cell(cell1, item, size_pt=9, is_html_like=True)
                    
                    cell2 = table_trc_word.cell(i, 1)
                    format_table_cell(cell2, value, size_pt=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                p_franch_title = doc_word.add_paragraph()
                set_paragraph_format(p_franch_title, space_before_pt=10, space_after_pt=10,
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1.18)
                run_franch_title = p_franch_title.add_run("MONTANTS DE FRANCHISES (par sinistre exprimés en €)")
                set_run_font(run_franch_title, name='Calibri', size_pt=11, bold=True)

                table_data_franchises_word = [
                    ("Dommages subis par les ouvrages de bâtiment", "XXXXXXXXXXX"), 
                    ("Catastrophes naturelles", "Montant déterminé par la loi ou par ses textes d'application"),
                    ("Responsabilité civile (1)", ""), 
                    ("- Assuré maître d'ouvrage", "XXXXXXXXXXX"),
                    ("- Assurés intervenants", "SANS"),
                    ("Maintenance-visite", "XXXXXXXXXXX")
                ]

                table_franchises_doc = doc_word.add_table(rows=len(table_data_franchises_word), cols=2)
                table_franchises_doc.style = 'TableGrid'
                table_franchises_doc.autofit = True

                for i, (item, value) in enumerate(table_data_franchises_word):
                    cell1 = table_franchises_doc.cell(i, 0)
                    cell2 = table_franchises_doc.cell(i, 1)
                    
                    indent = 0
                    if item.startswith("- "): 
                        indent = 0.27

                    format_table_cell(cell1, item, size_pt=9, left_indent_inches=indent)
                    
                    if item == "Responsabilité civile (1)":
                        cell1.merge(cell2) 
                    else:
                        format_table_cell(cell2, value, size_pt=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                p_footnote = doc_word.add_paragraph()
                set_paragraph_format(p_footnote, space_before_pt=6, space_after_pt=0, 
                                     line_spacing_rule=WD_LINE_SPACING.SINGLE) 
                run_footnote = p_footnote.add_run("     (1) Ces franchises s'appliquent pour des dommages autres que corporels")
                set_run_font(run_footnote, name='Calibri', size_pt=10)


                p_spacer_after_footnote = doc_word.add_paragraph()
                set_paragraph_format(p_spacer_after_footnote, space_before_pt=0, space_after_pt=1)

                p_date = doc_word.add_paragraph()
                set_paragraph_format(p_date, space_before_pt=5, space_after_pt=5,
                                     line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_value=1)
                run_date = p_date.add_run(f"Date de simulation de tarif: le {date_generation}")
                set_run_font(run_date, name='Calibri', size_pt=11)

                p_reserve = doc_word.add_paragraph()
                set_paragraph_format(p_reserve, space_before_pt=6, space_after_pt=1,
                                     line_spacing_rule=WD_LINE_SPACING.SINGLE)
                run_reserve = p_reserve.add_run("(*) Cette tarification est faite sous réserve d'acceptation du risque par la compagnie")
                set_run_font(run_reserve, name='Calibri', size_pt=10)

                doc_word.save(word_buffer)
                word_value = word_buffer.getvalue()
                word_buffer.close()
                response = HttpResponse(word_value, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{filename_base}.docx"'
                return response
            except Exception as e:
                return Response({'error': f"Erreur lors de la génération du document Word: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        else:
            return Response({'error': 'Type de document non supporté.'}, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['get'], url_path='history')
    def history(self, request, pk=None):
        proposal = get_object_or_404(Proposal, pk=pk)
        history_entries = ProposalHistory.objects.filter(proposal=proposal).order_by('-timestamp')
        serializer = ProposalHistorySerializer(history_entries, many=True)
        return Response(serializer.data)

